"""
Spacing Optimizer for IOST BSc CSIT Internship Reports
Optimizes spacing between sections, subsections, headings, and content
according to IOST Tribhuvan University guidelines
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re


class SpacingOptimizer:
    """Optimize spacing according to IOST CSIT standards"""

    def __init__(self, document_path):
        self.doc = Document(document_path)

        # IOST CSIT Spacing Standards (based on your requirements)
        self.spacing_standards = {
            "line_spacing": 1.5,  # 1.5 for all paragraphs
            "chapter_space_before": Pt(12),  # Space before chapter headings
            "chapter_space_after": Pt(12),  # Space after chapter headings
            "section_space_before": Pt(6),  # Space before section headings
            "section_space_after": Pt(6),  # Space after section headings
            "subsection_space_before": Pt(6),  # Space before subsection headings
            "subsection_space_after": Pt(6),  # Space after subsection headings
            "paragraph_space_before": Pt(0),  # No space before regular paragraphs
            "paragraph_space_after": Pt(0),  # No space after regular paragraphs
            "caption_space_before": Pt(6),  # Space before figure/table captions
            "caption_space_after": Pt(6),  # Space after figure/table captions
        }

    def is_chapter_heading(self, text):
        """Check if text is a chapter heading"""
        patterns = [r"^Chapter\s+\d+", r"^CHAPTER\s+\d+", r"^chapter\s+\d+"]
        return any(re.match(pattern, text.strip(), re.IGNORECASE) for pattern in patterns)

    def is_section_heading(self, text):
        """Check if text is a section heading (e.g., 1.1, 2.3)"""
        pattern = r"^\d+\.\d+\s"
        return bool(re.match(pattern, text.strip()))

    def is_subsection_heading(self, text):
        """Check if text is a subsection heading (e.g., 1.1.1, 2.3.4)"""
        pattern = r"^\d+\.\d+\.\d+\s"
        return bool(re.match(pattern, text.strip()))

    def is_figure_caption(self, text):
        """Check if text is a figure caption"""
        patterns = [r"^Figure\s+\d+\.\d+", r"^FIGURE\s+\d+\.\d+", r"^figure\s+\d+\.\d+"]
        return any(re.match(pattern, text.strip(), re.IGNORECASE) for pattern in patterns)

    def is_table_caption(self, text):
        """Check if text is a table caption"""
        patterns = [r"^Table\s+\d+\.\d+", r"^TABLE\s+\d+\.\d+", r"^table\s+\d+\.\d+"]
        return any(re.match(pattern, text.strip(), re.IGNORECASE) for pattern in patterns)

    def analyze_current_spacing(self):
        """Analyze current spacing issues in the document"""
        print("🔍 ANALYZING CURRENT SPACING")
        print("=" * 40)

        spacing_issues = []

        for i, paragraph in enumerate(self.doc.paragraphs):
            if not paragraph.text.strip():
                continue

            text = paragraph.text.strip()
            para_format = paragraph.paragraph_format

            # Check line spacing
            if para_format.line_spacing != self.spacing_standards["line_spacing"]:
                spacing_issues.append(
                    {
                        "paragraph": i + 1,
                        "type": "line_spacing",
                        "text_preview": text[:50] + "..." if len(text) > 50 else text,
                        "current": para_format.line_spacing,
                        "expected": self.spacing_standards["line_spacing"],
                    }
                )

            # Check space before/after based on heading type
            if self.is_chapter_heading(text):
                if para_format.space_before != self.spacing_standards["chapter_space_before"]:
                    spacing_issues.append(
                        {
                            "paragraph": i + 1,
                            "type": "chapter_space_before",
                            "text_preview": text[:50] + "..." if len(text) > 50 else text,
                            "current": para_format.space_before,
                            "expected": self.spacing_standards["chapter_space_before"],
                        }
                    )

                if para_format.space_after != self.spacing_standards["chapter_space_after"]:
                    spacing_issues.append(
                        {
                            "paragraph": i + 1,
                            "type": "chapter_space_after",
                            "text_preview": text[:50] + "..." if len(text) > 50 else text,
                            "current": para_format.space_after,
                            "expected": self.spacing_standards["chapter_space_after"],
                        }
                    )

            elif self.is_section_heading(text):
                if para_format.space_before != self.spacing_standards["section_space_before"]:
                    spacing_issues.append(
                        {
                            "paragraph": i + 1,
                            "type": "section_space_before",
                            "text_preview": text[:50] + "..." if len(text) > 50 else text,
                            "current": para_format.space_before,
                            "expected": self.spacing_standards["section_space_before"],
                        }
                    )

            elif self.is_subsection_heading(text):
                if para_format.space_before != self.spacing_standards["subsection_space_before"]:
                    spacing_issues.append(
                        {
                            "paragraph": i + 1,
                            "type": "subsection_space_before",
                            "text_preview": text[:50] + "..." if len(text) > 50 else text,
                            "current": para_format.space_before,
                            "expected": self.spacing_standards["subsection_space_before"],
                        }
                    )

        return spacing_issues

    def optimize_paragraph_spacing(self, paragraph, text):
        """Optimize spacing for a specific paragraph based on its type"""
        para_format = paragraph.paragraph_format

        # Set line spacing to 1.5 for all paragraphs
        para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if self.is_chapter_heading(text):
            # Chapter headings: 12pt before and after
            para_format.space_before = self.spacing_standards["chapter_space_before"]
            para_format.space_after = self.spacing_standards["chapter_space_after"]

        elif self.is_section_heading(text):
            # Section headings: 6pt before and after
            para_format.space_before = self.spacing_standards["section_space_before"]
            para_format.space_after = self.spacing_standards["section_space_after"]

        elif self.is_subsection_heading(text):
            # Subsection headings: 6pt before and after
            para_format.space_before = self.spacing_standards["subsection_space_before"]
            para_format.space_after = self.spacing_standards["subsection_space_after"]

        elif self.is_figure_caption(text) or self.is_table_caption(text):
            # Figure/Table captions: 6pt before and after
            para_format.space_before = self.spacing_standards["caption_space_before"]
            para_format.space_after = self.spacing_standards["caption_space_after"]

        else:
            # Regular paragraphs: no extra spacing
            para_format.space_before = self.spacing_standards["paragraph_space_before"]
            para_format.space_after = self.spacing_standards["paragraph_space_after"]

    def remove_excessive_empty_paragraphs(self):
        """Remove excessive empty paragraphs that create unwanted spacing"""
        print("🧹 Removing excessive empty paragraphs...")

        paragraphs_to_remove = []
        consecutive_empty = 0

        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip() == "" and len(paragraph.runs) == 0:
                consecutive_empty += 1
                # Allow maximum 1 empty paragraph between sections
                if consecutive_empty > 1:
                    paragraphs_to_remove.append(paragraph)
            else:
                consecutive_empty = 0

        # Remove excessive empty paragraphs
        removed_count = 0
        for para in paragraphs_to_remove:
            try:
                p = para._element
                p.getparent().remove(p)
                removed_count += 1
            except:
                pass

        print(f"   Removed {removed_count} excessive empty paragraphs")
        return removed_count

    def optimize_spacing_between_elements(self):
        """Optimize spacing between different document elements"""
        print("📏 Optimizing spacing between elements...")

        optimized_count = 0

        for i, paragraph in enumerate(self.doc.paragraphs):
            if not paragraph.text.strip():
                continue

            text = paragraph.text.strip()

            # Optimize spacing for this paragraph
            self.optimize_paragraph_spacing(paragraph, text)
            optimized_count += 1

        print(f"   Optimized spacing for {optimized_count} paragraphs")
        return optimized_count

    def fix_heading_content_spacing(self):
        """Ensure proper spacing between headings and their content"""
        print("🔗 Fixing heading-content spacing...")

        fixed_count = 0
        paragraphs = list(self.doc.paragraphs)

        for i, paragraph in enumerate(paragraphs):
            if not paragraph.text.strip():
                continue

            text = paragraph.text.strip()

            # If this is a heading, check the next paragraph
            if (
                self.is_chapter_heading(text)
                or self.is_section_heading(text)
                or self.is_subsection_heading(text)
            ):

                # Look for the next non-empty paragraph (content)
                for j in range(i + 1, min(i + 3, len(paragraphs))):
                    next_para = paragraphs[j]
                    if next_para.text.strip():
                        # Ensure content paragraph has no extra space before
                        next_para.paragraph_format.space_before = Pt(0)
                        fixed_count += 1
                        break

        print(f"   Fixed spacing for {fixed_count} heading-content pairs")
        return fixed_count

    def optimize_all_spacing(self):
        """Main method to optimize all spacing in the document"""
        print("🎯 OPTIMIZING DOCUMENT SPACING")
        print("=" * 40)

        # Step 1: Analyze current issues
        issues = self.analyze_current_spacing()
        print(f"📊 Found {len(issues)} spacing issues")

        # Step 2: Remove excessive empty paragraphs
        removed = self.remove_excessive_empty_paragraphs()

        # Step 3: Optimize spacing between elements
        optimized = self.optimize_spacing_between_elements()

        # Step 4: Fix heading-content spacing
        fixed = self.fix_heading_content_spacing()

        print(f"\n✅ SPACING OPTIMIZATION COMPLETED!")
        print(f"   📊 Issues found: {len(issues)}")
        print(f"   🧹 Empty paragraphs removed: {removed}")
        print(f"   📏 Paragraphs optimized: {optimized}")
        print(f"   🔗 Heading-content pairs fixed: {fixed}")

        return {
            "issues_found": len(issues),
            "empty_removed": removed,
            "paragraphs_optimized": optimized,
            "heading_content_fixed": fixed,
            "detailed_issues": issues,
        }

    def generate_spacing_report(self):
        """Generate a detailed spacing analysis report"""
        print("\n📋 DETAILED SPACING ANALYSIS REPORT")
        print("=" * 50)

        issues = self.analyze_current_spacing()

        if not issues:
            print("✅ No spacing issues found!")
            return

        # Group issues by type
        issue_types = {}
        for issue in issues:
            issue_type = issue["type"]
            if issue_type not in issue_types:
                issue_types[issue_type] = []
            issue_types[issue_type].append(issue)

        for issue_type, type_issues in issue_types.items():
            print(f"\n📌 {issue_type.replace('_', ' ').title()}: {len(type_issues)} issues")
            for i, issue in enumerate(type_issues[:5], 1):  # Show first 5
                print(f"   {i}. Paragraph {issue['paragraph']}: {issue['text_preview']}")
                print(f"      Current: {issue['current']} → Expected: {issue['expected']}")

            if len(type_issues) > 5:
                print(f"   ... and {len(type_issues) - 5} more")

    def save_document(self, output_path):
        """Save the spacing-optimized document"""
        self.doc.save(output_path)
        print(f"💾 Spacing-optimized document saved as: {output_path}")
        return output_path


def optimize_document_spacing(input_path, output_path=None):
    """Main function to optimize spacing in a document"""
    if output_path is None:
        base_name = input_path.replace(".docx", "")
        output_path = f"spacing_optimized_{base_name}.docx"

    print(f"🚀 Starting spacing optimization for: {input_path}")

    try:
        optimizer = SpacingOptimizer(input_path)

        # Generate initial report
        optimizer.generate_spacing_report()

        # Optimize spacing
        results = optimizer.optimize_all_spacing()

        # Save optimized document
        optimizer.save_document(output_path)

        print(f"\n🎉 SUCCESS! Spacing-optimized document created: {output_path}")

        return output_path, results

    except Exception as e:
        print(f"❌ Error optimizing spacing: {str(e)}")
        return None, None


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python spacing_optimizer.py <input_file.docx> [output_file.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    optimize_document_spacing(input_file, output_file)
