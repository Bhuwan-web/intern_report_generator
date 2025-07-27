#!/usr/bin/env python3
"""
Test script to check if all dependencies are working correctly
"""

import sys
import os


def test_imports():
    """Test if all required modules can be imported"""
    print("🔍 Testing imports...")

    try:
        import docx

        print("✅ python-docx imported successfully")
    except ImportError as e:
        print(f"❌ python-docx import failed: {e}")
        print("   Fix: pip install python-docx")
        return False

    try:
        import re

        print("✅ re module imported successfully")
    except ImportError as e:
        print(f"❌ re module import failed: {e}")
        return False

    try:
        from document_formatter import IOSTDocumentFormatter

        print("✅ document_formatter imported successfully")
    except ImportError as e:
        print(f"❌ document_formatter import failed: {e}")
        print("   Make sure document_formatter.py is in the same directory")
        return False

    try:
        from academic_tools import analyze_document

        print("✅ academic_tools imported successfully")
    except ImportError as e:
        print(f"❌ academic_tools import failed: {e}")
        print("   Make sure academic_tools.py is in the same directory")
        return False

    try:
        from enhanced_academic_tools import auto_fix_document

        print("✅ enhanced_academic_tools imported successfully")
    except ImportError as e:
        print(f"❌ enhanced_academic_tools import failed: {e}")
        print("   Make sure enhanced_academic_tools.py is in the same directory")
        return False

    return True


def test_file_access():
    """Test if test files exist and are accessible"""
    print("\n📁 Testing file access...")

    # Import docx here after we've tested it exists
    try:
        import docx
    except ImportError:
        print("❌ Cannot import docx for file testing")
        return False

    test_files = ["my_report.docx"]

    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"✅ Found: {file_path}")
            try:
                doc = docx.Document(file_path)
                print(f"✅ Can read: {file_path}")
            except Exception as e:
                print(f"❌ Cannot read {file_path}: {e}")
                return False
        else:
            print(f"⚠️  Not found: {file_path} (this is okay if you're using a different file)")

    return True


def test_basic_functionality():
    """Test basic functionality with a simple document"""
    print("\n🧪 Testing basic functionality...")

    try:
        # Import docx here
        import docx

        # Create a simple test document
        doc = docx.Document()
        doc.add_paragraph("Chapter 1: Introduction")
        doc.add_paragraph("This is a test paragraph with I implemented something.")
        doc.add_paragraph("Visit https://python.org for more information.")

        test_file = "test_document.docx"
        doc.save(test_file)
        print(f"✅ Created test document: {test_file}")

        # Test formatter
        from document_formatter import IOSTDocumentFormatter

        formatter = IOSTDocumentFormatter(test_file)
        formatter.format_document()
        formatter.save_document("test_formatted.docx")
        print("✅ Document formatting test passed")

        # Test auto-fixer
        from enhanced_academic_tools import DocumentAutoFixer

        fixer = DocumentAutoFixer(test_file)
        results = fixer.fix_document(create_backup=False)
        fixer.save_fixed_document("test_fixed.docx")
        print("✅ Auto-fix test passed")

        # Clean up test files
        for f in [test_file, "test_formatted.docx", "test_fixed.docx"]:
            if os.path.exists(f):
                os.remove(f)

        print("✅ All basic functionality tests passed")
        return True

    except Exception as e:
        print(f"❌ Basic functionality test failed: {e}")
        return False


def main():
    print("IOST Document Formatter - Setup Test")
    print("=" * 50)

    all_tests_passed = True

    # Test imports
    if not test_imports():
        all_tests_passed = False

    # Test file access
    if not test_file_access():
        all_tests_passed = False

    # Test basic functionality
    if not test_basic_functionality():
        all_tests_passed = False

    print("\n" + "=" * 50)
    if all_tests_passed:
        print("🎉 ALL TESTS PASSED!")
        print("You can now use the formatter with:")
        print("  python main.py my_report.docx")
    else:
        print("❌ SOME TESTS FAILED!")
        print("Please fix the issues above before using the formatter")

    print("\n📋 Common fixes:")
    print("1. Install python-docx: pip install python-docx")
    print("2. Make sure all .py files are in the same directory")
    print("3. Make sure your .docx file exists and is not corrupted")


if __name__ == "__main__":
    main()
