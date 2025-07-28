#!/usr/bin/env python3
"""
Test runner for IOST BSc CSIT Internship Report Formatter
Runs all test suites and provides comprehensive testing results
"""

import sys
import os
import subprocess
from pathlib import Path

# Add src directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))


def run_test_module(test_module_path, test_name):
    """Run a specific test module and return results"""
    print(f"\n{'='*60}")
    print(f"🧪 Running {test_name}")
    print(f"{'='*60}")

    try:
        result = subprocess.run(
            [sys.executable, test_module_path],
            capture_output=True,
            text=True,
            cwd=os.path.dirname(os.path.dirname(__file__)),  # Run from project root
        )

        if result.returncode == 0:
            print(f"✅ {test_name} PASSED")
            print(result.stdout)
            return True
        else:
            print(f"❌ {test_name} FAILED")
            print("STDOUT:", result.stdout)
            print("STDERR:", result.stderr)
            return False

    except Exception as e:
        print(f"❌ {test_name} ERROR: {e}")
        return False


def run_all_tests():
    """Run all test suites"""
    print("🚀 IOST Academic Tools - Comprehensive Test Suite")
    print("=" * 70)

    test_dir = Path(__file__).parent
    test_results = []

    # Define test modules
    test_modules = [
        ("test_document_formatter_fix.py", "Document Formatter Fix Tests"),
        ("test_line_breaks.py", "Academic Line Break Manager Tests"),
        ("test_section_breaks.py", "Section Break Preservation Tests"),
        ("test_page_breaks.py", "Page Break Management Tests"),
        ("test_spacing.py", "Spacing Optimization Tests"),
    ]

    # Run each test module
    for test_file, test_name in test_modules:
        test_path = test_dir / test_file
        if test_path.exists():
            success = run_test_module(str(test_path), test_name)
            test_results.append((test_name, success))
        else:
            print(f"⚠️  Test file not found: {test_file}")
            test_results.append((test_name, False))

    # Summary
    print(f"\n{'='*70}")
    print("📊 TEST SUMMARY")
    print(f"{'='*70}")

    passed = sum(1 for _, success in test_results if success)
    total = len(test_results)

    for test_name, success in test_results:
        status = "✅ PASSED" if success else "❌ FAILED"
        print(f"   {status} - {test_name}")

    print(f"\n🎯 Results: {passed}/{total} tests passed")

    if passed == total:
        print("🎉 All tests passed! The system is working correctly.")
        return True
    else:
        print("⚠️  Some tests failed. Please review the output above.")
        return False


def run_integration_tests():
    """Run integration tests with the main pipeline"""
    print(f"\n{'='*60}")
    print("🔗 Integration Tests")
    print(f"{'='*60}")

    # Test main.py with different actions
    test_actions = ["format", "analyze", "autofix", "linebreaks"]

    # Create a simple test document
    from docx import Document

    doc = Document()
    doc.add_paragraph("Chapter 1: Introduction")
    doc.add_paragraph("This is a test document for integration testing.")
    doc.add_paragraph("1. Literature Review")
    doc.add_paragraph("Literature review content.")

    test_file = "integration_test.docx"
    doc.save(test_file)

    integration_results = []

    try:
        for action in test_actions:
            print(f"\n🧪 Testing main.py with action: {action}")

            try:
                result = subprocess.run(
                    [sys.executable, "src/main.py", test_file, action],
                    capture_output=True,
                    text=True,
                    timeout=60,  # 60 second timeout
                    cwd=".",
                )

                if result.returncode == 0:
                    print(f"   ✅ {action} action completed successfully")
                    integration_results.append((action, True))
                else:
                    print(f"   ❌ {action} action failed")
                    print(f"   Error: {result.stderr}")
                    integration_results.append((action, False))

            except subprocess.TimeoutExpired:
                print(f"   ⏰ {action} action timed out")
                integration_results.append((action, False))
            except Exception as e:
                print(f"   ❌ {action} action error: {e}")
                integration_results.append((action, False))

    finally:
        # Clean up test files
        cleanup_patterns = [
            test_file,
            f"formatted_{test_file.replace('.docx', '')}.docx",
            f"auto_fixed_{test_file.replace('.docx', '')}.docx",
            f"line_break_fixed_{test_file.replace('.docx', '')}.docx",
            f"final_{test_file.replace('.docx', '')}.docx",
            f"backup_{test_file}",
        ]

        for pattern in cleanup_patterns:
            if os.path.exists(pattern):
                os.remove(pattern)

    # Integration test summary
    passed_integration = sum(1 for _, success in integration_results if success)
    total_integration = len(integration_results)

    print(f"\n📊 Integration Test Results: {passed_integration}/{total_integration} passed")

    for action, success in integration_results:
        status = "✅ PASSED" if success else "❌ FAILED"
        print(f"   {status} - {action} action")

    return passed_integration == total_integration


def main():
    """Main test runner"""
    print("🎯 Starting Comprehensive Test Suite...")

    # Run unit tests
    unit_tests_passed = run_all_tests()

    # Run integration tests
    integration_tests_passed = run_integration_tests()

    # Final summary
    print(f"\n{'='*70}")
    print("🏁 FINAL TEST RESULTS")
    print(f"{'='*70}")

    if unit_tests_passed and integration_tests_passed:
        print("🎉 ALL TESTS PASSED!")
        print("   ✅ Unit tests: PASSED")
        print("   ✅ Integration tests: PASSED")
        print("\n🚀 The IOST Academic Tools system is ready for use!")
        return 0
    else:
        print("❌ SOME TESTS FAILED")
        print(
            f"   {'✅' if unit_tests_passed else '❌'} Unit tests: {'PASSED' if unit_tests_passed else 'FAILED'}"
        )
        print(
            f"   {'✅' if integration_tests_passed else '❌'} Integration tests: {'PASSED' if integration_tests_passed else 'FAILED'}"
        )
        print("\n⚠️  Please review the test output and fix any issues.")
        return 1


if __name__ == "__main__":
    sys.exit(main())
