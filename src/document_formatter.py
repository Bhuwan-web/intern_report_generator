"""
Document Formatter for IOST BSc CSIT Internship Reports
Formats Word documents according to Tribhuvan University guidelines
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import argparse
import os


class IOSTDocumentFormatter:
    def fix_heading_line_breaks(self):
        """Ensure proper line breaks between headings and paragraphs, and remove excess empty lines"""
        paragraphs = self.doc.paragraphs
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i]
            text = para.text.strip()

            # Check for headings
            is_chapter = self.is_chapter_heading(text)
            is_section = any(
                text.startswith(f"{x}.{y}") for x in range(1, 20) for y in range(1, 20)
            ) and not any(
                text.startswith(f"{x}.{y}.{z}")
                for x in range(1, 20)
                for y in range(1, 20)
                for z in range(1, 20)
            )
            is_subsection = any(
                text.startswith(f"{x}.{y}.{z}")
                for x in range(1, 20)
                for y in range(1, 20)
                for z in range(1, 20)
            )

            # Remove empty paragraphs before headings
            if (is_chapter or is_section or is_subsection) and i > 0:
                prev_para = paragraphs[i - 1]
                if prev_para.text.strip() == "":
                    p = prev_para._element
                    try:
                        p.getparent().remove(p)
                        # After removal, stay at same index
                        i -= 1
                        continue
                    except Exception:
                        pass

                # Skip automatic line break insertion - let the academic line break manager handle this
                # This prevents errors and delegates proper spacing to the specialized line break manager
                # Remove extra empty paragraphs after heading
                j = i + 2
                while j < len(paragraphs) and paragraphs[j].text.strip() == "":
                    p = paragraphs[j]._element
                    try:
                        p.getparent().remove(p)
                    except Exception:
                        pass
                    j += 1

            i += 1

    def __init__(self, document_path):
        """Initialize formatter with document path"""
        self.doc_path = document_path
        self.doc = Document(document_path)
        self.setup_styles()

    def setup_styles(self):
        """Setup custom styles according to IOST guidelines"""
        styles = self.doc.styles

        # Normal paragraph style
        if "IOST Normal" not in [s.name for s in styles]:
            normal_style = styles.add_style("IOST Normal", WD_STYLE_TYPE.PARAGRAPH)
            normal_font = normal_style.font
            normal_font.name = "Times New Roman"
            normal_font.size = Pt(12)

            normal_paragraph = normal_style.paragraph_format
            normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            normal_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            normal_paragraph.space_after = Pt(0)
            normal_paragraph.space_before = Pt(0)

        # Chapter heading style (16pt, Bold)
        if "IOST Chapter" not in [s.name for s in styles]:
            chapter_style = styles.add_style("IOST Chapter", WD_STYLE_TYPE.PARAGRAPH)
            chapter_font = chapter_style.font
            chapter_font.name = "Times New Roman"
            chapter_font.size = Pt(16)
            chapter_font.bold = True

            chapter_paragraph = chapter_style.paragraph_format
            chapter_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            chapter_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            chapter_paragraph.space_after = Pt(12)
            chapter_paragraph.space_before = Pt(12)

        # Section heading style (14pt, Bold)
        if "IOST Section" not in [s.name for s in styles]:
            section_style = styles.add_style("IOST Section", WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = "Times New Roman"
            section_font.size = Pt(14)
            section_font.bold = True

            section_paragraph = section_style.paragraph_format
            section_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            section_paragraph.space_after = Pt(6)
            section_paragraph.space_before = Pt(6)

        # Sub-section heading style (12pt, Bold)
        if "IOST Subsection" not in [s.name for s in styles]:
            subsection_style = styles.add_style("IOST Subsection", WD_STYLE_TYPE.PARAGRAPH)
            subsection_font = subsection_style.font
            subsection_font.name = "Times New Roman"
            subsection_font.size = Pt(12)
            subsection_font.bold = True

            subsection_paragraph = subsection_style.paragraph_format
            subsection_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            subsection_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            subsection_paragraph.space_after = Pt(6)
            subsection_paragraph.space_before = Pt(6)

        # Figure/Table caption style (12pt, Bold, Centered)
        if "IOST Caption" not in [s.name for s in styles]:
            caption_style = styles.add_style("IOST Caption", WD_STYLE_TYPE.PARAGRAPH)
            caption_font = caption_style.font
            caption_font.name = "Times New Roman"
            caption_font.size = Pt(12)
            caption_font.bold = True

            caption_paragraph = caption_style.paragraph_format
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            caption_paragraph.space_after = Pt(6)
            caption_paragraph.space_before = Pt(6)

    def set_page_margins(self):
        """Set page margins according to IOST guidelines"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.0)
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height

    def format_paragraphs(self):
        """Format all paragraphs according to IOST guidelines"""
        for paragraph in self.doc.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue

            # Determine paragraph type and apply appropriate style
            text = paragraph.text.strip()

            # Check if it's a chapter heading (starts with "Chapter" or contains chapter-like patterns)
            if (
                text.startswith("Chapter ")
                or text.startswith("CHAPTER ")
                or any(text.startswith(f"{i}.") for i in range(1, 20))
            ):
                paragraph.style = self.doc.styles["IOST Chapter"]

            # Check if it's a section heading (numbered like 1.1, 2.3, etc.)
            elif any(
                text.startswith(f"{i}.{j}") for i in range(1, 20) for j in range(1, 20)
            ) and not any(
                text.startswith(f"{i}.{j}.{k}")
                for i in range(1, 20)
                for j in range(1, 20)
                for k in range(1, 20)
            ):
                paragraph.style = self.doc.styles["IOST Section"]

            # Check if it's a subsection heading (numbered like 1.1.1, 2.3.4, etc.)
            elif any(
                text.startswith(f"{i}.{j}.{k}")
                for i in range(1, 20)
                for j in range(1, 20)
                for k in range(1, 20)
            ):
                paragraph.style = self.doc.styles["IOST Subsection"]

            # Check if it's a figure or table caption
            elif (
                text.startswith("Figure ")
                or text.startswith("Table ")
                or text.startswith("FIGURE ")
                or text.startswith("TABLE ")
            ):
                paragraph.style = self.doc.styles["IOST Caption"]

            # Default to normal paragraph style
            else:
                paragraph.style = self.doc.styles["IOST Normal"]

    def format_tables(self):
        """Format tables according to IOST guidelines"""
        for table in self.doc.tables:
            # Center align the table
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Format table cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Set font for table content
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)

                        # Set paragraph formatting
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_page_numbers(self):
        """Add page numbering (Note: This is basic - manual adjustment may be needed)"""
        # This is a simplified implementation
        # For complex page numbering (Roman/Arabic), manual adjustment in Word is recommended
        pass

    def is_chapter_heading(self, text):
        """Check if text is a chapter heading"""
        return (
            text.startswith("Chapter ")
            or text.startswith("CHAPTER ")
            or text.startswith("chapter ")
        )

    def clean_page_breaks(self):
        """Remove unnecessary page breaks and section breaks"""
        print("Cleaning unnecessary page breaks...")

        # Find paragraphs with manual page breaks
        paragraphs_to_check = []

        for i, paragraph in enumerate(self.doc.paragraphs):
            # Check for manual page breaks in empty paragraphs
            if paragraph.text.strip() == "":
                for run in paragraph.runs:
                    # Check if run contains a page break
                    if hasattr(run._element, "xml") and "w:br" in str(run._element.xml):
                        paragraphs_to_check.append((i, paragraph))
                        break

        # Remove unnecessary page breaks (keep only those before chapters)
        for i, paragraph in paragraphs_to_check:
            # Check if next non-empty paragraph is a chapter
            next_chapter = False
            for j in range(i + 1, len(self.doc.paragraphs)):
                next_para = self.doc.paragraphs[j]
                if next_para.text.strip():
                    if self.is_chapter_heading(next_para.text.strip()):
                        next_chapter = True
                    break

            # Remove page break if it's not before a chapter
            if not next_chapter:
                p = paragraph._element
                p.getparent().remove(p)

    def add_chapter_page_breaks(self):
        """Add proper page breaks before chapters"""
        print("Adding proper chapter page breaks...")

        # Process paragraphs in reverse order to avoid index issues when inserting
        paragraphs = list(self.doc.paragraphs)

        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()

            if self.is_chapter_heading(text):
                # Skip page break for first chapter or if it's at the beginning
                if i == 0:
                    continue

                # Check if this is Chapter 1 (don't add page break before first chapter)
                if any(
                    text.lower().startswith(prefix)
                    for prefix in ["chapter 1", "chapter 1:", "chapter 1."]
                ):
                    continue

                # Check if there's already a page break before this chapter
                has_page_break = False

                # Look at previous paragraphs for existing page breaks
                for j in range(max(0, i - 5), i):
                    if j < len(paragraphs):
                        prev_para = paragraphs[j]
                        if prev_para.text.strip() == "":
                            for run in prev_para.runs:
                                if hasattr(run._element, "xml") and "w:br" in str(
                                    run._element.xml
                                ):
                                    has_page_break = True
                                    break
                        if has_page_break:
                            break

                # Add page break if none exists
                if not has_page_break:
                    # Create a new paragraph with page break
                    new_para = paragraph.insert_paragraph_before()
                    run = new_para.runs[0] if new_para.runs else new_para.add_run()
                    from docx.enum.text import WD_BREAK

                    run.add_break(WD_BREAK.PAGE)  # Page break

    def remove_excessive_empty_paragraphs(self):
        """Remove excessive empty paragraphs (keep max 1 between sections)"""
        print("Removing excessive empty paragraphs...")

        paragraphs_to_remove = []
        empty_count = 0

        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip() == "":
                empty_count += 1
                # Keep only first empty paragraph in a sequence, remove others
                if empty_count > 1:
                    paragraphs_to_remove.append(paragraph)
            else:
                empty_count = 0

        # Remove excessive empty paragraphs
        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)

    def format_document(self):
        """Apply all formatting rules to the document"""
        print("Setting page margins...")
        self.set_page_margins()

        print("Cleaning unnecessary page breaks...")
        self.clean_page_breaks()

        print("Removing excessive empty paragraphs...")
        self.remove_excessive_empty_paragraphs()

        print("Adding proper chapter page breaks...")
        self.add_chapter_page_breaks()

        print("Formatting paragraphs...")
        self.format_paragraphs()

        print("Fixing heading line breaks...")
        self.fix_heading_line_breaks()

        print("Formatting tables...")
        self.format_tables()

        print("Applying IOST spacing standards...")
        self.apply_iost_spacing()

        print("Document formatting completed!")

    def apply_iost_spacing(self):
        """Apply IOST spacing standards to all paragraphs"""
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING

        for paragraph in self.doc.paragraphs:
            if not paragraph.text.strip():
                continue

            text = paragraph.text.strip()
            para_format = paragraph.paragraph_format

            # Set 1.5 line spacing for all paragraphs
            para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # Apply spacing based on paragraph type
            if text.startswith("Chapter ") or text.startswith("CHAPTER "):
                # Chapter headings: 12pt before and after
                para_format.space_before = Pt(12)
                para_format.space_after = Pt(12)

            elif any(text.startswith(f"{i}.{j} ") for i in range(1, 20) for j in range(1, 20)):
                # Section headings: 6pt before and after
                para_format.space_before = Pt(6)
                para_format.space_after = Pt(6)

            elif any(
                text.startswith(f"{i}.{j}.{k} ")
                for i in range(1, 20)
                for j in range(1, 20)
                for k in range(1, 20)
            ):
                # Subsection headings: 6pt before and after
                para_format.space_before = Pt(6)
                para_format.space_after = Pt(6)

            elif (
                text.startswith("Figure ")
                or text.startswith("Table ")
                or text.startswith("FIGURE ")
                or text.startswith("TABLE ")
            ):
                # Figure/Table captions: 6pt before and after
                para_format.space_before = Pt(6)
                para_format.space_after = Pt(6)

            else:
                # Regular paragraphs: no extra spacing
                para_format.space_before = Pt(0)
                para_format.space_after = Pt(0)

    def save_document(self, output_path=None):
        """Save the formatted document"""
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(self.doc_path))[0]
            output_path = f"formatted_{base_name}.docx"

        self.doc.save(output_path)
        print(f"Formatted document saved as: {output_path}")
        return output_path


def main():
    """Main function to run the formatter"""
    parser = argparse.ArgumentParser(
        description="Format Word documents according to IOST BSc CSIT guidelines"
    )
    parser.add_argument("input_file", help="Path to the input Word document")
    parser.add_argument("-o", "--output", help="Output file path (optional)")

    args = parser.parse_args()

    if not os.path.exists(args.input_file):
        print(f"Error: File '{args.input_file}' not found!")
        return

    try:
        print(f"Processing document: {args.input_file}")
        formatter = IOSTDocumentFormatter(args.input_file)
        formatter.format_document()
        output_file = formatter.save_document(args.output)

        print("\n" + "=" * 50)
        print("FORMATTING COMPLETED SUCCESSFULLY!")
        print("=" * 50)
        print(f"Input file: {args.input_file}")
        print(f"Output file: {output_file}")
        print("\nNote: For complete compliance, please manually verify:")
        print("- Page numbering (Roman numerals for preliminary pages)")
        print("- Chapter page breaks and numbering")
        print("- Figure and table positioning")
        print("- Citation formatting (consider using Zotero or Mendeley)")

    except Exception as e:
        print(f"Error processing document: {str(e)}")
        print("Make sure the input file is a valid Word document (.docx)")


if __name__ == "__main__":
    main()
