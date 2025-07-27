"""
Academic Tools for IOST BSc CSIT Internship Reports
Includes language checking, citation validation, and formatting assistance
"""

import re
import sys
from typing import List, Dict, Tuple
from docx import Document
import argparse


class LanguageChecker:
    """Check for common language issues in academic writing"""

    def __init__(self):
        # First-person pronouns to avoid (except in preliminary sections)
        self.first_person_pronouns = [
            r"\bI\b",
            r"\bwe\b",
            r"\byou\b",
            r"\bme\b",
            r"\bus\b",
            r"\bmy\b",
            r"\bour\b",
            r"\byour\b",
            r"\bmyself\b",
            r"\bourselves\b",
            r"\byourself\b",
        ]

        # Common passive voice indicators (preferred in academic writing)
        self.active_voice_indicators = [
            r"\bI\s+\w+ed\b",
            r"\bwe\s+\w+ed\b",
            r"\bI\s+will\b",
            r"\bwe\s+will\b",
            r"\bI\s+have\b",
            r"\bwe\s+have\b",
            r"\bI\s+did\b",
            r"\bwe\s+did\b",
        ]

        # Common contractions to avoid in formal writing
        self.contractions = [
            r"don't",
            r"won't",
            r"can't",
            r"shouldn't",
            r"wouldn't",
            r"couldn't",
            r"isn't",
            r"aren't",
            r"wasn't",
            r"weren't",
            r"haven't",
            r"hasn't",
            r"hadn't",
            r"doesn't",
            r"didn't",
            r"it's",
            r"that's",
            r"there's",
        ]

    def check_first_person(self, text: str) -> List[Dict]:
        """Check for first-person pronouns"""
        issues = []
        for pronoun in self.first_person_pronouns:
            matches = re.finditer(pronoun, text, re.IGNORECASE)
            for match in matches:
                issues.append(
                    {
                        "type": "first_person",
                        "text": match.group(),
                        "position": match.span(),
                        "suggestion": "Consider using passive voice or third person",
                    }
                )
        return issues

    def check_contractions(self, text: str) -> List[Dict]:
        """Check for contractions"""
        issues = []
        for contraction in self.contractions:
            matches = re.finditer(contraction, text, re.IGNORECASE)
            for match in matches:
                expanded = self.expand_contraction(match.group().lower())
                issues.append(
                    {
                        "type": "contraction",
                        "text": match.group(),
                        "position": match.span(),
                        "suggestion": f'Use "{expanded}" instead',
                    }
                )
        return issues

    def expand_contraction(self, contraction: str) -> str:
        """Expand common contractions"""
        expansions = {
            "don't": "do not",
            "won't": "will not",
            "can't": "cannot",
            "shouldn't": "should not",
            "wouldn't": "would not",
            "couldn't": "could not",
            "isn't": "is not",
            "aren't": "are not",
            "wasn't": "was not",
            "weren't": "were not",
            "haven't": "have not",
            "hasn't": "has not",
            "hadn't": "had not",
            "doesn't": "does not",
            "didn't": "did not",
            "it's": "it is",
            "that's": "that is",
            "there's": "there is",
        }
        return expansions.get(contraction, contraction)

    def check_document(self, doc_path: str) -> Dict:
        """Check entire document for language issues"""
        doc = Document(doc_path)
        all_issues = []

        for para_num, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if not text.strip():
                continue

            # Skip preliminary sections (basic heuristic)
            if any(
                keyword in text.lower()
                for keyword in ["certificate", "acknowledgment", "abstract"]
            ):
                continue

            para_issues = []
            para_issues.extend(self.check_first_person(text))
            para_issues.extend(self.check_contractions(text))

            if para_issues:
                all_issues.append(
                    {
                        "paragraph": para_num + 1,
                        "text_preview": text[:100] + "..." if len(text) > 100 else text,
                        "issues": para_issues,
                    }
                )

        return {
            "total_issues": sum(len(para["issues"]) for para in all_issues),
            "paragraphs_with_issues": all_issues,
        }


class CitationChecker:
    """Check for proper APA citation format"""

    def __init__(self):
        # Basic APA in-text citation patterns
        self.apa_patterns = [
            r"\([A-Za-z]+,\s*\d{4}\)",  # (Author, Year)
            r"\([A-Za-z]+\s+&\s+[A-Za-z]+,\s*\d{4}\)",  # (Author & Author, Year)
            r"\([A-Za-z]+\s+et\s+al\.,\s*\d{4}\)",  # (Author et al., Year)
        ]

        # Common citation issues
        self.url_pattern = r"https?://[^\s]+"
        self.bare_url_pattern = r"(?<![\(\[])\bhttps?://[^\s\)\]]+(?![\)\]])"

    def find_citations(self, text: str) -> List[Dict]:
        """Find potential citations in text"""
        citations = []

        for pattern in self.apa_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                citations.append(
                    {
                        "type": "apa_citation",
                        "text": match.group(),
                        "position": match.span(),
                        "status": "valid",
                    }
                )

        return citations

    def find_bare_urls(self, text: str) -> List[Dict]:
        """Find URLs that should be properly cited"""
        issues = []
        matches = re.finditer(self.bare_url_pattern, text)

        for match in matches:
            issues.append(
                {
                    "type": "bare_url",
                    "text": match.group(),
                    "position": match.span(),
                    "suggestion": "URLs should be properly cited in APA format or moved to references",
                }
            )

        return issues

    def check_document(self, doc_path: str) -> Dict:
        """Check document for citation issues"""
        doc = Document(doc_path)
        citations_found = []
        citation_issues = []

        for para_num, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if not text.strip():
                continue

            # Find valid citations
            para_citations = self.find_citations(text)
            if para_citations:
                citations_found.extend(para_citations)

            # Find citation issues
            bare_urls = self.find_bare_urls(text)
            if bare_urls:
                citation_issues.append(
                    {
                        "paragraph": para_num + 1,
                        "text_preview": text[:100] + "..." if len(text) > 100 else text,
                        "issues": bare_urls,
                    }
                )

        return {
            "citations_found": len(citations_found),
            "citation_details": citations_found,
            "citation_issues": citation_issues,
            "total_issues": sum(len(para["issues"]) for para in citation_issues),
        }


def analyze_document(doc_path: str):
    """Comprehensive document analysis"""
    print(f"Analyzing document: {doc_path}")
    print("=" * 60)

    # Language check
    print("\n📝 LANGUAGE ANALYSIS")
    print("-" * 30)
    lang_checker = LanguageChecker()
    lang_results = lang_checker.check_document(doc_path)

    if lang_results["total_issues"] == 0:
        print("✅ No language issues found!")
    else:
        print(f"⚠️  Found {lang_results['total_issues']} language issues:")
        for para in lang_results["paragraphs_with_issues"]:
            print(f"\nParagraph {para['paragraph']}:")
            print(f"Preview: {para['text_preview']}")
            for issue in para["issues"]:
                print(f"  - {issue['type'].title()}: '{issue['text']}' → {issue['suggestion']}")

    # Citation check
    print("\n📚 CITATION ANALYSIS")
    print("-" * 30)
    citation_checker = CitationChecker()
    citation_results = citation_checker.check_document(doc_path)

    print(f"✅ Found {citation_results['citations_found']} properly formatted citations")

    if citation_results["total_issues"] == 0:
        print("✅ No citation issues found!")
    else:
        print(f"⚠️  Found {citation_results['total_issues']} citation issues:")
        for para in citation_results["citation_issues"]:
            print(f"\nParagraph {para['paragraph']}:")
            print(f"Preview: {para['text_preview']}")
            for issue in para["issues"]:
                print(f"  - {issue['type'].title()}: '{issue['text']}' → {issue['suggestion']}")

    # Recommendations
    print("\n🔧 RECOMMENDATIONS")
    print("-" * 30)
    print("For better academic writing:")
    print("• Use Grammarly or LanguageTool for grammar checking")
    print("• Use Zotero, Mendeley, or EndNote for citation management")
    print("• Consider Hemingway Editor for readability")
    print("• Use passive voice in technical descriptions")
    print("• Ensure all sources are properly cited in APA format")


def main():
    parser = argparse.ArgumentParser(description="Academic writing analysis for IOST reports")
    parser.add_argument("document", help="Path to the Word document")

    args = parser.parse_args()

    if not args.document.endswith(".docx"):
        print("Error: Please provide a .docx file")
        return

    try:
        analyze_document(args.document)
    except Exception as e:
        print(f"Error analyzing document: {str(e)}")


if __name__ == "__main__":
    main()
