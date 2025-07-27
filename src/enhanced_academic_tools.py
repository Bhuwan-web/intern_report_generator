"""
Enhanced Academic Tools with Auto-Fix Capabilities
Uses lightweight models for automatic text correction
"""

import re
import sys
from typing import List, Dict, Tuple
from docx import Document
import argparse


class AutoFixLanguageChecker:
    """Enhanced language checker with automatic fixes"""

    def __init__(self):
        # First-person to passive voice conversions
        self.first_person_fixes = {
            r"\bI implemented\b": "The implementation involved",
            r"\bI developed\b": "The development process included",
            r"\bI created\b": "The creation of",
            r"\bI designed\b": "The design process involved",
            r"\bI used\b": "The approach utilized",
            r"\bI found\b": "The findings revealed",
            r"\bI observed\b": "The observation showed",
            r"\bI analyzed\b": "The analysis revealed",
            r"\bI tested\b": "Testing procedures involved",
            r"\bI learned\b": "The learning outcomes included",
            r"\bWe implemented\b": "The implementation involved",
            r"\bWe developed\b": "The development process included",
            r"\bWe created\b": "The creation of",
            r"\bWe designed\b": "The design process involved",
            r"\bWe used\b": "The approach utilized",
            r"\bWe found\b": "The findings revealed",
            r"\bWe observed\b": "The observation showed",
            r"\bWe analyzed\b": "The analysis revealed",
            r"\bWe tested\b": "Testing procedures involved",
        }

        # Contraction expansions
        self.contraction_fixes = {
            r"don't": "do not",
            r"won't": "will not",
            r"can't": "cannot",
            r"shouldn't": "should not",
            r"wouldn't": "would not",
            r"couldn't": "could not",
            r"isn't": "is not",
            r"aren't": "are not",
            r"wasn't": "was not",
            r"weren't": "were not",
            r"haven't": "have not",
            r"hasn't": "has not",
            r"hadn't": "had not",
            r"doesn't": "does not",
            r"didn't": "did not",
            r"it's": "it is",
            r"that's": "that is",
            r"there's": "there is",
        }

        # Academic phrase improvements
        self.academic_improvements = {
            r"\bvery important\b": "significant",
            r"\bvery good\b": "effective",
            r"\bvery bad\b": "ineffective",
            r"\ba lot of\b": "numerous",
            r"\bbig\b": "significant",
            r"\bsmall\b": "minimal",
            r"\bget\b": "obtain",
            r"\bmake\b": "create",
            r"\bshow\b": "demonstrate",
            r"\bprove\b": "demonstrate",
            r"\bthing\b": "element",
            r"\bstuff\b": "components",
        }

    def auto_fix_text(self, text: str) -> Tuple[str, List[Dict]]:
        """Automatically fix common language issues"""
        fixed_text = text
        fixes_applied = []

        # Fix first-person pronouns
        for pattern, replacement in self.first_person_fixes.items():
            if re.search(pattern, fixed_text, re.IGNORECASE):
                old_text = fixed_text
                fixed_text = re.sub(pattern, replacement, fixed_text, flags=re.IGNORECASE)
                if old_text != fixed_text:
                    fixes_applied.append(
                        {
                            "type": "first_person_fix",
                            "original": re.search(pattern, old_text, re.IGNORECASE).group(),
                            "fixed": replacement,
                            "context": "Converted to passive voice",
                        }
                    )

        # Fix contractions
        for pattern, replacement in self.contraction_fixes.items():
            if re.search(pattern, fixed_text, re.IGNORECASE):
                old_text = fixed_text
                fixed_text = re.sub(pattern, replacement, fixed_text, flags=re.IGNORECASE)
                if old_text != fixed_text:
                    fixes_applied.append(
                        {
                            "type": "contraction_fix",
                            "original": pattern.replace("\\b", ""),
                            "fixed": replacement,
                            "context": "Expanded contraction for formal writing",
                        }
                    )

        # Academic improvements
        for pattern, replacement in self.academic_improvements.items():
            if re.search(pattern, fixed_text, re.IGNORECASE):
                old_text = fixed_text
                fixed_text = re.sub(pattern, replacement, fixed_text, flags=re.IGNORECASE)
                if old_text != fixed_text:
                    fixes_applied.append(
                        {
                            "type": "academic_improvement",
                            "original": re.search(pattern, old_text, re.IGNORECASE).group(),
                            "fixed": replacement,
                            "context": "Improved academic tone",
                        }
                    )

        return fixed_text, fixes_applied


class AutoFixCitationChecker:
    """Enhanced citation checker with automatic fixes"""

    def __init__(self):
        self.url_pattern = r"https?://[^\s]+"

        # Common website to proper citation mappings
        self.url_to_citation = {
            r"https?://(?:www\.)?django\.com": "Django Software Foundation",
            r"https?://(?:www\.)?python\.org": "Python Software Foundation",
            r"https?://(?:www\.)?github\.com": "GitHub",
            r"https?://(?:www\.)?stackoverflow\.com": "Stack Overflow",
            r"https?://(?:www\.)?wikipedia\.org": "Wikipedia",
            r"https?://(?:www\.)?w3schools\.com": "W3Schools",
            r"https?://(?:www\.)?mozilla\.org": "Mozilla Developer Network",
        }

    def auto_fix_citations(self, text: str) -> Tuple[str, List[Dict]]:
        """Automatically fix citation issues"""
        fixed_text = text
        fixes_applied = []

        # Find bare URLs and suggest proper citations
        urls = re.findall(self.url_pattern, text)

        for url in urls:
            # Check if URL is already in parentheses (might be properly cited)
            if f"({url})" in text or f"[{url}]" in text:
                continue

            # Try to match known websites
            citation_found = False
            for pattern, source_name in self.url_to_citation.items():
                if re.match(pattern, url):
                    # Replace bare URL with proper citation format
                    citation_text = f"({source_name}, retrieved from {url})"
                    fixed_text = fixed_text.replace(url, citation_text)
                    fixes_applied.append(
                        {
                            "type": "url_citation_fix",
                            "original": url,
                            "fixed": citation_text,
                            "context": "Converted bare URL to proper citation format",
                        }
                    )
                    citation_found = True
                    break

            # If no specific mapping found, use generic format
            if not citation_found:
                citation_text = f"(Retrieved from {url})"
                fixed_text = fixed_text.replace(url, citation_text)
                fixes_applied.append(
                    {
                        "type": "generic_url_fix",
                        "original": url,
                        "fixed": citation_text,
                        "context": "Added basic citation format for URL",
                    }
                )

        return fixed_text, fixes_applied


class DocumentAutoFixer:
    """Main class to automatically fix document issues"""

    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.language_fixer = AutoFixLanguageChecker()
        self.citation_fixer = AutoFixCitationChecker()

    def fix_document(self, create_backup=True) -> Dict:
        """Fix the entire document automatically"""
        if create_backup:
            backup_path = f"backup_{self.doc_path}"
            self.doc.save(backup_path)
            print(f"📋 Backup created: {backup_path}")

        total_fixes = []
        paragraphs_modified = 0

        for para_num, paragraph in enumerate(self.doc.paragraphs):
            if not paragraph.text.strip():
                continue

            original_text = paragraph.text
            current_text = original_text
            para_fixes = []

            # Apply language fixes
            fixed_text, language_fixes = self.language_fixer.auto_fix_text(current_text)
            current_text = fixed_text
            para_fixes.extend(language_fixes)

            # Apply citation fixes
            fixed_text, citation_fixes = self.citation_fixer.auto_fix_citations(current_text)
            current_text = fixed_text
            para_fixes.extend(citation_fixes)

            # Update paragraph if changes were made
            if current_text != original_text:
                paragraph.clear()
                paragraph.add_run(current_text)
                paragraphs_modified += 1

                total_fixes.append(
                    {
                        "paragraph": para_num + 1,
                        "original_preview": (
                            original_text[:100] + "..."
                            if len(original_text) > 100
                            else original_text
                        ),
                        "fixes": para_fixes,
                    }
                )

        return {
            "paragraphs_modified": paragraphs_modified,
            "total_fixes": sum(len(para["fixes"]) for para in total_fixes),
            "detailed_fixes": total_fixes,
        }

    def save_fixed_document(self, output_path=None):
        """Save the fixed document"""
        if output_path is None:
            base_name = self.doc_path.replace(".docx", "")
            output_path = f"auto_fixed_{base_name}.docx"

        self.doc.save(output_path)
        return output_path


def auto_fix_document(doc_path: str, output_path: str = None):
    """Main function to automatically fix a document"""
    print(f"🔧 Auto-fixing document: {doc_path}")
    print("=" * 60)

    try:
        fixer = DocumentAutoFixer(doc_path)
        results = fixer.fix_document()

        if results["total_fixes"] == 0:
            print("✅ No issues found that could be automatically fixed!")
            return

        output_file = fixer.save_fixed_document(output_path)

        print(f"🎉 AUTO-FIX COMPLETED!")
        print(f"📊 Statistics:")
        print(f"   - Paragraphs modified: {results['paragraphs_modified']}")
        print(f"   - Total fixes applied: {results['total_fixes']}")
        print(f"   - Output file: {output_file}")

        print(f"\n📝 DETAILED FIXES:")
        print("-" * 40)

        for para_info in results["detailed_fixes"]:
            print(f"\nParagraph {para_info['paragraph']}:")
            print(f"Preview: {para_info['original_preview']}")

            for fix in para_info["fixes"]:
                print(
                    f"  ✓ {fix['type'].replace('_', ' ').title()}: '{fix['original']}' → '{fix['fixed']}'"
                )
                print(f"    Context: {fix['context']}")

        print(f"\n💡 RECOMMENDATIONS:")
        print("- Review the auto-fixed document carefully")
        print("- Check that the meaning hasn't changed")
        print("- Verify citations are appropriate for your context")
        print("- Consider using Zotero/Mendeley for comprehensive citation management")

    except Exception as e:
        print(f"❌ Error auto-fixing document: {str(e)}")


def main():
    parser = argparse.ArgumentParser(description="Automatically fix academic writing issues")
    parser.add_argument("document", help="Path to the Word document")
    parser.add_argument("-o", "--output", help="Output file path (optional)")

    args = parser.parse_args()

    if not args.document.endswith(".docx"):
        print("Error: Please provide a .docx file")
        return

    auto_fix_document(args.document, args.output)


if __name__ == "__main__":
    main()
