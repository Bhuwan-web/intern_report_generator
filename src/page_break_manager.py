"""
Advanced Page Break Manager for IOST Documents
Handles complex page break scenarios and section management
"""

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import qn
import re


class PageBreakManager:
    """Advanced page break management for academic documents"""

    def __init__(self, document_path):
        self.doc = Document(document_path)
        self.chapter_patterns = [
            r"^Chapter\s+\d+",
            r"^CHAPTER\s+\d+",
            r"^chapter\s+\d+",
            r"^\d+\.\s*Introduction",
            r"^\d+\.\s*Conclusion",
        ]

    def is_chapter_heading(self, text):
        """Enhanced chapter detection"""
        text = text.strip()
        for pattern in self.chapter_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False

    def is_preliminary_section(self, text):
        """Check if text is a preliminary section"""
        preliminary_sections = [
            "certificate",
            "acknowledgment",
            "acknowledgement",
            "abstract",
            "table of contents",
            "list of figures",
            "list of tables",
            "list of abbreviations",
            "executive summary",
        ]
        return any(section in text.lower() for section in preliminary_sections)

    def find_all_breaks(self):
        """Find all existing page breaks and section breaks"""
        breaks_info = []

        for i, paragraph in enumerate(self.doc.paragraphs):
            # Check for page breaks in runs
            for run in paragraph.runs:
                if hasattr(run._element, "xml"):
                    xml_str = str(run._element.xml)
                    if "w:br" in xml_str and 'type="page"' in xml_str:
                        breaks_info.append(
                            {
                                "type": "page_break",
                                "paragraph_index": i,
                                "paragraph": paragraph,
                                "text_preview": (
                                    paragraph.text[:50] if paragraph.text else "[Empty]"
                                ),
                            }
                        )

        # Check for section breaks
        for i, section in enumerate(self.doc.sections):
            if i > 0:  # Skip first section
                breaks_info.append(
                    {
                        "type": "section_break",
                        "section_index": i,
                        "section": section,
                        "start_type": section.start_type,
                    }
                )

        return breaks_info

    def remove_unnecessary_breaks(self):
        """Remove page breaks that aren't before chapters or major sections"""
        print("🧹 Removing unnecessary page breaks...")

        paragraphs_to_remove = []

        for i, paragraph in enumerate(self.doc.paragraphs):
            # Check if paragraph contains only a page break
            if paragraph.text.strip() == "":
                has_page_break = False
                for run in paragraph.runs:
                    if hasattr(run._element, "xml"):
                        xml_str = str(run._element.xml)
                        if "w:br" in xml_str and 'type="page"' in xml_str:
                            has_page_break = True
                            break

                if has_page_break:
                    # Check what comes after this page break
                    next_significant_para = None
                    for j in range(i + 1, min(i + 5, len(self.doc.paragraphs))):
                        if self.doc.paragraphs[j].text.strip():
                            next_significant_para = self.doc.paragraphs[j]
                            break

                    # Keep page break only if it's before a chapter or major section
                    should_keep = False
                    if next_significant_para:
                        text = next_significant_para.text.strip()
                        if self.is_chapter_heading(text) or self.is_preliminary_section(text):
                            should_keep = True

                    if not should_keep:
                        paragraphs_to_remove.append(paragraph)

        # Remove unnecessary page break paragraphs
        removed_count = 0
        for para in paragraphs_to_remove:
            try:
                p = para._element
                p.getparent().remove(p)
                removed_count += 1
            except:
                pass  # Skip if already removed or error

        print(f"   Removed {removed_count} unnecessary page breaks")

    def add_chapter_page_breaks(self):
        """Add page breaks before all chapters (except first)"""
        print("📄 Adding proper chapter page breaks...")

        chapters_found = []

        # Find all chapter headings
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()
            if self.is_chapter_heading(text):
                chapters_found.append((i, paragraph, text))

        print(f"   Found {len(chapters_found)} chapters")

        # Add page breaks before chapters (except first)
        added_breaks = 0

        for i, (para_index, paragraph, text) in enumerate(chapters_found):
            # Skip first chapter
            if i == 0:
                continue

            # Check if there's already a page break before this chapter
            has_page_break = False

            # Look back up to 5 paragraphs for existing page break
            for j in range(max(0, para_index - 5), para_index):
                check_para = self.doc.paragraphs[j]
                if check_para.text.strip() == "":
                    for run in check_para.runs:
                        if hasattr(run._element, "xml"):
                            xml_str = str(run._element.xml)
                            if "w:br" in xml_str and 'type="page"' in xml_str:
                                has_page_break = True
                                break
                if has_page_break:
                    break

            # Add page break if none exists
            if not has_page_break:
                try:
                    new_para = paragraph.insert_paragraph_before()
                    run = new_para.add_run()
                    run.add_break(break_type=6)  # Page break
                    added_breaks += 1
                except:
                    print(f"   Warning: Could not add page break before: {text[:50]}")

        print(f"   Added {added_breaks} chapter page breaks")

    def clean_empty_paragraphs(self):
        """Remove excessive empty paragraphs"""
        print("🧽 Cleaning excessive empty paragraphs...")

        paragraphs_to_remove = []
        consecutive_empty = 0

        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip() == "" and len(paragraph.runs) == 0:
                consecutive_empty += 1
                # Keep only first empty paragraph in sequence
                if consecutive_empty > 1:
                    paragraphs_to_remove.append(paragraph)
            else:
                consecutive_empty = 0

        # Remove excessive empty paragraphs
        removed_count = 0
        for para in paragraphs_to_remove:
            try:
                p = para._element
                p.getparent().remove(p)
                removed_count += 1
            except:
                pass

        print(f"   Removed {removed_count} excessive empty paragraphs")

    def optimize_page_breaks(self):
        """Main method to optimize all page breaks"""
        print("🔧 Optimizing page breaks and document structure...")

        # Step 1: Remove unnecessary breaks
        self.remove_unnecessary_breaks()

        # Step 2: Clean empty paragraphs
        self.clean_empty_paragraphs()

        # Step 3: Add proper chapter breaks
        self.add_chapter_page_breaks()

        print("✅ Page break optimization completed!")

    def generate_break_report(self):
        """Generate a report of all breaks in the document"""
        breaks = self.find_all_breaks()

        print("\n📊 PAGE BREAK ANALYSIS REPORT")
        print("=" * 50)

        page_breaks = [b for b in breaks if b["type"] == "page_break"]
        section_breaks = [b for b in breaks if b["type"] == "section_break"]

        print(f"📄 Page breaks found: {len(page_breaks)}")
        for i, pb in enumerate(page_breaks, 1):
            print(f"   {i}. Paragraph {pb['paragraph_index']}: {pb['text_preview']}")

        print(f"\n📑 Section breaks found: {len(section_breaks)}")
        for i, sb in enumerate(section_breaks, 1):
            print(f"   {i}. Section {sb['section_index']}: {sb['start_type']}")

        # Find chapters
        chapters = []
        for i, paragraph in enumerate(self.doc.paragraphs):
            if self.is_chapter_heading(paragraph.text.strip()):
                chapters.append((i, paragraph.text.strip()))

        print(f"\n📚 Chapters found: {len(chapters)}")
        for i, (para_idx, text) in enumerate(chapters, 1):
            print(f"   {i}. Paragraph {para_idx}: {text}")

    def save_document(self, output_path):
        """Save the optimized document"""
        self.doc.save(output_path)
        print(f"💾 Optimized document saved as: {output_path}")
        return output_path


def optimize_document_breaks(input_path, output_path=None):
    """Main function to optimize page breaks in a document"""
    if output_path is None:
        base_name = input_path.replace(".docx", "")
        output_path = f"optimized_{base_name}.docx"

    print(f"🚀 Starting page break optimization for: {input_path}")

    try:
        manager = PageBreakManager(input_path)

        # Generate initial report
        manager.generate_break_report()

        # Optimize breaks
        manager.optimize_page_breaks()

        # Save optimized document
        manager.save_document(output_path)

        print(f"\n🎉 SUCCESS! Optimized document created: {output_path}")

        return output_path

    except Exception as e:
        print(f"❌ Error optimizing document: {str(e)}")
        return None


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python page_break_manager.py <input_file.docx> [output_file.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    optimize_document_breaks(input_file, output_file)
