"""
Academic Line Break Manager
Implements proper line break standards for academic documents following APA/MLA guidelines
"""

import re
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import argparse


class AcademicLineBreakManager:
    """Manages line breaks according to academic formatting standards"""

    def __init__(self):
        # Define standard line break requirements
        self.line_break_standards = {
            "chapter_heading": {
                "before": 2,  # 2 line breaks before chapter
                "after": 1,  # 1 line break after chapter
                "description": 'Chapter headings (e.g., "Chapter 1", "CHAPTER 1")',
            },
            "major_section": {
                "before": 2,  # 2 line breaks before major sections
                "after": 1,  # 1 line break after major sections
                "description": 'Major sections (e.g., "1. Introduction", "2. Literature Review")',
            },
            "section_heading": {
                "before": 1,  # 1 line break before section headings
                "after": 0,  # No extra line break after section headings
                "description": 'Section headings (e.g., "1.1 Background", "2.1 Overview")',
            },
            "subsection_heading": {
                "before": 1,  # 1 line break before subsection headings
                "after": 0,  # No extra line break after subsection headings
                "description": 'Subsection headings (e.g., "1.1.1 Problem Statement")',
            },
            "paragraph": {
                "before": 0,  # No extra line break before paragraphs
                "after": 0,  # No extra line break after paragraphs
                "description": "Regular paragraphs",
            },
            "figure_table": {
                "before": 1,  # 1 line break before figures/tables
                "after": 1,  # 1 line break after figures/tables
                "description": "Figures and tables with captions",
            },
            "list_item": {
                "before": 0,  # No extra line break before list items
                "after": 0,  # No extra line break after list items
                "description": "Bulleted or numbered list items",
            },
            "quote_block": {
                "before": 1,  # 1 line break before block quotes
                "after": 1,  # 1 line break after block quotes
                "description": "Block quotations",
            },
        }

        # Patterns to identify different content types
        self.content_patterns = {
            "chapter_heading": [
                r"^Chapter\s+\d+",
                r"^CHAPTER\s+\d+",
                r"^chapter\s+\d+",
            ],
            "major_section": [
                r"^\d+\.\s+[A-Z][^.]*$",  # "1. Introduction"
                r"^\d+\s+[A-Z][^.]*$",  # "1 Introduction"
            ],
            "section_heading": [
                r"^\d+\.\d+\s+[A-Z][^.]*$",  # "1.1 Background"
                r"^\d+\.\d+\.\s+[A-Z][^.]*$",  # "1.1. Background"
            ],
            "subsection_heading": [
                r"^\d+\.\d+\.\d+\s+[A-Z][^.]*$",  # "1.1.1 Problem Statement"
                r"^\d+\.\d+\.\d+\.\s+[A-Z][^.]*$",  # "1.1.1. Problem Statement"
            ],
            "figure_table": [
                r"^Figure\s+\d+",
                r"^Table\s+\d+",
                r"^FIGURE\s+\d+",
                r"^TABLE\s+\d+",
            ],
            "list_item": [
                r"^\s*[-•]\s+",  # Bulleted lists
                r"^\s*\d+\.\s+",  # Numbered lists (but not headings)
                r"^\s*[a-z]\)\s+",  # Lettered lists
            ],
            "quote_block": [
                r'^".*"$',  # Quoted text
                r"^\s{4,}",  # Indented text (potential quote)
            ],
        }

    def identify_content_type(self, text: str) -> str:
        """Identify the type of content based on text patterns"""
        text = text.strip()

        if not text:
            return "empty"

        # Check each content type
        for content_type, patterns in self.content_patterns.items():
            for pattern in patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    # Additional validation for list items vs headings
                    if content_type == "list_item" and re.match(r"^\d+\.\s+[A-Z][^.]*$", text):
                        # This might be a major section, not a list item
                        if len(text.split()) <= 5:  # Short titles are likely headings
                            continue
                    return content_type

        return "paragraph"

    def analyze_document_structure(self, doc: Document) -> List[Dict]:
        """Analyze the document structure and identify content types"""
        structure = []

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            content_type = self.identify_content_type(text)

            structure.append(
                {
                    "index": i,
                    "paragraph": paragraph,
                    "text": text,
                    "content_type": content_type,
                    "current_spacing_before": self._get_spacing_before(paragraph),
                    "current_spacing_after": self._get_spacing_after(paragraph),
                    "required_spacing_before": self.line_break_standards.get(content_type, {}).get(
                        "before", 0
                    ),
                    "required_spacing_after": self.line_break_standards.get(content_type, {}).get(
                        "after", 0
                    ),
                }
            )

        return structure

    def _get_spacing_before(self, paragraph) -> float:
        """Get current spacing before paragraph in points"""
        try:
            if paragraph.paragraph_format.space_before:
                return float(paragraph.paragraph_format.space_before.pt)
        except:
            pass
        return 0.0

    def _get_spacing_after(self, paragraph) -> float:
        """Get current spacing after paragraph in points"""
        try:
            if paragraph.paragraph_format.space_after:
                return float(paragraph.paragraph_format.space_after.pt)
        except:
            pass
        return 0.0

    def preserve_section_breaks(self, doc: Document) -> Dict:
        """Identify and preserve existing section breaks from Google Docs imports"""
        print("� PRESYERVING EXISTING SECTION BREAKS")
        print("-" * 40)
        
        preserved_breaks = []
        
        try:
            # Analyze sections for page break types
            for i, section in enumerate(doc.sections):
                try:
                    if hasattr(section, '_sectPr') and section._sectPr is not None:
                        sectPr_xml = str(section._sectPr.xml)
                        
                        # Check for different types of section breaks
                        if 'nextPage' in sectPr_xml:
                            preserved_breaks.append({
                                'section_index': i,
                                'break_type': 'nextPage',
                                'description': 'Section break (next page)'
                            })
                        elif 'oddPage' in sectPr_xml:
                            preserved_breaks.append({
                                'section_index': i,
                                'break_type': 'oddPage', 
                                'description': 'Section break (odd page)'
                            })
                        elif 'evenPage' in sectPr_xml:
                            preserved_breaks.append({
                                'section_index': i,
                                'break_type': 'evenPage',
                                'description': 'Section break (even page)'
                            })
                        elif 'continuous' in sectPr_xml:
                            preserved_breaks.append({
                                'section_index': i,
                                'break_type': 'continuous',
                                'description': 'Section break (continuous)'
                            })
                except Exception as e:
                    print(f"   ⚠️  Could not analyze section {i}: {e}")
                    continue
        
        except Exception as e:
            print(f"   ⚠️  Could not analyze sections: {e}")
        
        if preserved_breaks:
            print(f"   ✅ Found {len(preserved_breaks)} section breaks to preserve:")
            for break_info in preserved_breaks:
                print(f"      - Section {break_info['section_index']}: {break_info['description']}")
        else:
            print("   ✅ No section breaks found (or document has simple structure)")
        
        return {
            'preserved_breaks': preserved_breaks,
            'total_preserved': len(preserved_breaks)
        }

    def apply_standard_line_breaks(self, doc: Document) -> Dict:
        """Apply standard line breaks throughout the document while preserving section breaks"""
        print("📏 APPLYING ACADEMIC LINE BREAK STANDARDS")
        print("=" * 50)
        
        # First preserve existing section breaks
        section_info = self.preserve_section_breaks(doc)

        structure = self.analyze_document_structure(doc)
        corrections_made = []

        for item in structure:
            paragraph = item["paragraph"]
            content_type = item["content_type"]

            if content_type == "empty":
                continue

            # Check if this paragraph is after a section break
            has_section_break = self._has_section_break_before(doc, item["index"])

            # Get required spacing
            required_before = item["required_spacing_before"]
            required_after = item["required_spacing_after"]
            current_before = item["current_spacing_before"]
            current_after = item["current_spacing_after"]

            # If there's a section break before this paragraph, reduce spacing requirements
            # since the section break already provides visual separation
            if has_section_break and content_type in ["chapter_heading", "major_section"]:
                required_before = max(0, required_before - 1)  # Reduce by 1 line break
                print(f"   📄 Reduced spacing for {content_type} after section break: {item['text'][:40]}...")

            # Convert line breaks to points (assuming 12pt font, 1 line break = 18pt)
            required_before_pt = required_before * 18
            required_after_pt = required_after * 18

            changes_made = []

            # Adjust spacing before
            if abs(current_before - required_before_pt) > 1:  # Allow 1pt tolerance
                paragraph.paragraph_format.space_before = Pt(required_before_pt)
                changes_made.append(f"before: {current_before}pt → {required_before_pt}pt")

            # Adjust spacing after
            if abs(current_after - required_after_pt) > 1:  # Allow 1pt tolerance
                paragraph.paragraph_format.space_after = Pt(required_after_pt)
                changes_made.append(f"after: {current_after}pt → {required_after_pt}pt")

            if changes_made:
                corrections_made.append(
                    {
                        "paragraph_index": item["index"],
                        "content_type": content_type,
                        "text_preview": (
                            item["text"][:60] + "..." if len(item["text"]) > 60 else item["text"]
                        ),
                        "changes": changes_made,
                        "has_section_break": has_section_break
                    }
                )

                print(f"   ✅ Fixed {content_type}: {item['text'][:50]}...")
                for change in changes_made:
                    print(f"      → Spacing {change}")

        return {
            "total_corrections": len(corrections_made),
            "corrections_made": corrections_made,
            "document_structure": structure,
            "section_breaks_preserved": section_info["total_preserved"]
        }

    def remove_excessive_empty_paragraphs(self, doc: Document) -> int:
        """Remove excessive empty paragraphs while preserving required spacing"""
        print("\n🧹 REMOVING EXCESSIVE EMPTY PARAGRAPHS")
        print("-" * 40)

        paragraphs_to_remove = []
        consecutive_empty = 0

        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                consecutive_empty += 1
                # Remove if more than 2 consecutive empty paragraphs
                if consecutive_empty > 2:
                    paragraphs_to_remove.append(paragraph)
            else:
                consecutive_empty = 0

        # Remove excessive empty paragraphs
        removed_count = 0
        for para in paragraphs_to_remove:
            try:
                p = para._p
                p.getparent().remove(p)
                removed_count += 1
            except:
                pass

        if removed_count > 0:
            print(f"   ✅ Removed {removed_count} excessive empty paragraphs")
        else:
            print("   ✅ No excessive empty paragraphs found")

        return removed_count

    def add_required_page_breaks(self, doc: Document) -> int:
        """Add page breaks before chapters and major sections as needed"""
        print("\n📄 ADDING REQUIRED PAGE BREAKS")
        print("-" * 35)

        structure = self.analyze_document_structure(doc)
        page_breaks_added = 0

        for i, item in enumerate(structure):
            if item["content_type"] in ["chapter_heading", "major_section"]:
                # Skip first chapter/section
                if i == 0:
                    continue

                paragraph = item["paragraph"]

                # Check if there's already a page break before this
                has_page_break = self._has_page_break_before(doc, i)

                if not has_page_break:
                    # Add page break
                    try:
                        new_para = paragraph.insert_paragraph_before()
                        run = new_para.add_run()
                        from docx.enum.text import WD_BREAK

                        run.add_break(WD_BREAK.PAGE)
                        page_breaks_added += 1
                        print(f"   ✅ Added page break before: {item['text'][:50]}...")
                    except Exception as e:
                        print(f"   ⚠️  Could not add page break: {e}")

        if page_breaks_added == 0:
            print("   ✅ No additional page breaks needed")

        return page_breaks_added

    def _has_page_break_before(self, doc: Document, paragraph_index: int) -> bool:
        """Check if there's a page break or section break before the given paragraph"""
        # Look back up to 3 paragraphs for page breaks
        for i in range(max(0, paragraph_index - 3), paragraph_index):
            try:
                para = doc.paragraphs[i]
                if para.text.strip() == "":
                    for run in para.runs:
                        try:
                            if hasattr(run._element, "xml"):
                                xml_str = str(run._element.xml)
                                if "w:br" in xml_str and (
                                    'type="page"' in xml_str or 'type="nextPage"' in xml_str
                                ):
                                    return True
                        except:
                            continue
            except:
                continue

        # Also check for section breaks (from Google Docs imports)
        if self._has_section_break_before(doc, paragraph_index):
            return True

        return False

    def _has_section_break_before(self, doc: Document, paragraph_index: int) -> bool:
        """Check if there's a section break (next page) before the given paragraph"""
        try:
            # Check if this paragraph is at the start of a new section
            if paragraph_index == 0:
                return False

            # Get the current paragraph
            current_para = doc.paragraphs[paragraph_index]

            # Check if this paragraph starts a new section by examining sections
            current_section_index = None
            prev_section_index = None

            # Find which section this paragraph belongs to
            for section_idx, section in enumerate(doc.sections):
                try:
                    # Get the section's start element
                    section_start = section._sectPr.getparent()

                    # Check if current paragraph is after this section start
                    if self._is_paragraph_in_section(current_para, section):
                        current_section_index = section_idx
                        break
                except:
                    continue

            # If we found a section and it's not the first one, check for page break type
            if current_section_index is not None and current_section_index > 0:
                try:
                    section = doc.sections[current_section_index]
                    # Check section properties for page break
                    if hasattr(section, "_sectPr"):
                        sectPr_xml = str(section._sectPr.xml)
                        # Look for section break types that create new pages
                        if (
                            "nextPage" in sectPr_xml
                            or "oddPage" in sectPr_xml
                            or "evenPage" in sectPr_xml
                        ):
                            return True
                except:
                    pass

        except Exception:
            pass

        return False

    def _is_paragraph_in_section(self, paragraph, section) -> bool:
        """Check if a paragraph belongs to a specific section"""
        try:
            # This is a simplified check - in practice, section boundaries are complex
            # For now, we'll use a heuristic based on paragraph position
            para_element = paragraph._element
            section_element = section._sectPr.getparent() if hasattr(section, "_sectPr") else None

            if section_element is not None:
                # Check if paragraph comes after section start
                return True  # Simplified - assume yes for now
        except:
            pass
        return False

    def generate_line_break_report(self, structure: List[Dict]) -> str:
        """Generate a report of the document's line break structure"""
        report = []
        report.append("📊 DOCUMENT LINE BREAK ANALYSIS")
        report.append("=" * 50)

        # Count content types
        content_counts = {}
        for item in structure:
            content_type = item["content_type"]
            content_counts[content_type] = content_counts.get(content_type, 0) + 1

        report.append("\n📋 Content Structure:")
        for content_type, count in content_counts.items():
            if content_type != "empty":
                description = self.line_break_standards.get(content_type, {}).get(
                    "description", content_type
                )
                report.append(f"   - {description}: {count}")

        # Show line break standards
        report.append("\n📏 Applied Line Break Standards:")
        for content_type, standards in self.line_break_standards.items():
            if content_counts.get(content_type, 0) > 0:
                report.append(f"   - {standards['description']}:")
                report.append(f"     Before: {standards['before']} line breaks")
                report.append(f"     After: {standards['after']} line breaks")

        return "\n".join(report)


def process_document_line_breaks(input_path: str, output_path: str = None) -> str:
    """Process document line breaks according to academic standards"""
    if output_path is None:
        base_name = input_path.replace(".docx", "")
        output_path = f"line_break_fixed_{base_name}.docx"

    print("📏 ACADEMIC LINE BREAK PROCESSING")
    print("=" * 60)
    print(f"Processing: {input_path}")

    # Load document
    doc = Document(input_path)
    manager = AcademicLineBreakManager()

    # Step 1: Remove excessive empty paragraphs
    removed_empty = manager.remove_excessive_empty_paragraphs(doc)

    # Step 2: Apply standard line breaks
    line_break_results = manager.apply_standard_line_breaks(doc)

    # Step 3: Add required page breaks
    page_breaks_added = manager.add_required_page_breaks(doc)

    # Save processed document
    doc.save(output_path)

    # Generate and display report
    report = manager.generate_line_break_report(line_break_results["document_structure"])
    print(f"\n{report}")

    print(f"\n🎉 LINE BREAK PROCESSING COMPLETED!")
    print("=" * 50)
    print(f"   📝 Line break corrections: {line_break_results['total_corrections']}")
    print(f"   🧹 Empty paragraphs removed: {removed_empty}")
    print(f"   📄 Page breaks added: {page_breaks_added}")
    print(f"   💾 Saved as: {output_path}")

    return output_path


def main():
    """Main function for academic line break manager"""
    parser = argparse.ArgumentParser(description="Academic Line Break Manager")
    parser.add_argument("document", help="Path to the Word document")
    parser.add_argument("--output", help="Output file path")

    args = parser.parse_args()

    if not args.document.endswith(".docx"):
        print("Error: Please provide a .docx file")
        return

    try:
        process_document_line_breaks(args.document, args.output)
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback

        traceback.print_exc()


if __name__ == "__main__":
    main()
