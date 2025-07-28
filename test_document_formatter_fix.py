#!/usr/bin/env python3
"""
Test script to verify document formatter fixes
Tests that the insert_paragraph_after error is resolved
"""

import sys
import os
from docx import Document

# Add src directory to path
sys.path.append("src")


def create_test_document():
    """Create a test document with headings that would trigger the error"""
    doc = Document()

    # Add content that would trigger the line break insertion
    doc.add_paragraph("Chapter 1: Introduction")
    doc.add_paragraph("This paragraph immediately follows the chapter heading.")

    doc.add_paragraph("1. Literature Review")
    doc.add_paragraph("This section follows immediately.")

    doc.add_paragraph("1.1 Background")
    doc.add_paragraph("Background content here.")

    doc.add_paragraph("1.1.1 Problem Statement")
    doc.add_paragraph("Problem statement content.")

    test_file = "test_formatter_fix.docx"
    doc.save(test_file)
    return test_file


def test_document_formatter():
    """Test that document formatter works without errors"""
    print("🔧 Testing Document Formatter Fix")
    print("=" * 50)

    # Create test document
    test_file = create_test_document()
    print(f"✅ Created test document: {test_file}")

    try:
        from document_formatter import IOSTDocumentFormatter

        print("\n🔧 Testing Document Formatter:")
        print("-" * 40)

        # Initialize formatter
        formatter = IOSTDocumentFormatter(test_file)
        print("   ✅ Formatter initialized successfully")

        # Test formatting without errors
        formatter.format_document()
        print("   ✅ Document formatting completed without errors")

        # Save formatted document
        output_file = "test_formatter_fix_formatted.docx"
        formatter.save_document(output_file)
        print(f"   ✅ Formatted document saved: {output_file}")

        # Verify the output document can be opened
        formatted_doc = Document(output_file)
        print(f"   ✅ Output document verified ({len(formatted_doc.paragraphs)} paragraphs)")

        print(f"\n🎉 Document formatter test completed successfully!")

    except Exception as e:
        print(f"❌ Error during formatting: {e}")
        import traceback

        traceback.print_exc()
        return False

    finally:
        # Clean up test files
        for file in [test_file, "test_formatter_fix_formatted.docx"]:
            if os.path.exists(file):
                os.remove(file)
                print(f"🧹 Cleaned up: {file}")

    return True


def test_main_pipeline():
    """Test the main processing pipeline"""
    print("\n🔗 Testing Main Processing Pipeline:")
    print("-" * 50)

    # Create test document
    test_file = create_test_document()

    try:
        # Test using main.py
        import subprocess

        # Test the complete action
        result = subprocess.run(
            [sys.executable, "src/main.py", test_file, "format"],
            capture_output=True,
            text=True,
            cwd=".",
        )

        if result.returncode == 0:
            print("✅ Main pipeline completed successfully")
            print("Last few lines of output:")
            print(result.stdout[-200:])  # Show last 200 characters
        else:
            print("❌ Main pipeline failed")
            print("Error output:")
            print(result.stderr)
            return False

    except Exception as e:
        print(f"❌ Pipeline test failed: {e}")
        return False

    finally:
        # Clean up
        for file in [test_file, f"formatted_{test_file.replace('.docx', '')}.docx"]:
            if os.path.exists(file):
                os.remove(file)

    return True


def main():
    """Run the document formatter fix test"""
    print("🚀 Document Formatter Fix Test Suite")
    print("=" * 60)

    success = True

    # Test document formatter directly
    if not test_document_formatter():
        success = False

    # Test main pipeline
    if not test_main_pipeline():
        success = False

    print("\n" + "=" * 60)
    if success:
        print("🎉 All tests passed! Document formatter fix is working.")
    else:
        print("❌ Some tests failed. Check the output above.")

    print("\nThe fix addresses:")
    print("   ✅ Removed problematic insert_paragraph_after method")
    print("   ✅ Fixed break_type=6 issue with proper WD_BREAK.PAGE")
    print("   ✅ Delegated line break handling to academic line break manager")


if __name__ == "__main__":
    main()
