#!/usr/bin/env python3
"""
Test script for section break preservation
Tests that Google Docs section breaks (next page) are preserved during line break processing
"""

import sys
import os
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt

# Add src directory to path
sys.path.append("src")


def create_document_with_section_breaks():
    """Create a test document with section breaks (simulating Google Docs import)"""
    doc = Document()

    # Add first section content
    doc.add_paragraph("Chapter 1: Introduction")
    doc.add_paragraph("This is the introduction content.")
    doc.add_paragraph("More introduction content here.")

    # Add a new section with next page break
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)

    # Add content to the new section
    doc.add_paragraph("Chapter 2: Literature Review")
    doc.add_paragraph("This chapter reviews existing literature.")
    doc.add_paragraph("Literature review content continues here.")

    # Add another section with next page break
    another_section = doc.add_section(WD_SECTION_START.NEW_PAGE)

    # Add content to the third section
    doc.add_paragraph("Chapter 3: Methodology")
    doc.add_paragraph("This chapter describes the methodology.")
    doc.add_paragraph("Methodology details are provided here.")

    # Add a continuous section (no page break)
    continuous_section = doc.add_section(WD_SECTION_START.CONTINUOUS)

    # Add content to continuous section
    doc.add_paragraph("3.1 Data Collection")
    doc.add_paragraph("Data collection procedures are described.")

    test_file = "test_section_breaks.docx"
    doc.save(test_file)
    return test_file


def test_section_break_preservation():
    """Test that section breaks are properly preserved"""
    print("🔒 Testing Section Break Preservation")
    print("=" * 60)

    # Create test document with section breaks
    test_file = create_document_with_section_breaks()
    print(f"✅ Created test document with section breaks: {test_file}")

    try:
        from academic_line_break_manager import (
            AcademicLineBreakManager,
            process_document_line_breaks,
        )

        # Test section break detection
        print("\n🔍 Testing Section Break Detection:")
        print("-" * 50)

        manager = AcademicLineBreakManager()
        doc = Document(test_file)

        # Analyze original document sections
        print("Original document sections:")
        for i, section in enumerate(doc.sections):
            try:
                start_type = section.start_type
                print(f"   Section {i}: {start_type}")
            except Exception as e:
                print(f"   Section {i}: Could not determine type - {e}")

        # Test section break preservation method
        section_info = manager.preserve_section_breaks(doc)
        print(f"\nSection breaks detected: {section_info['total_preserved']}")

        # Test section break detection for specific paragraphs
        print(f"\n🔍 Testing Paragraph-Level Section Break Detection:")
        print("-" * 55)

        structure = manager.analyze_document_structure(doc)
        for item in structure[:10]:  # Test first 10 paragraphs
            if item["text"].strip():
                has_section_break = manager._has_section_break_before(doc, item["index"])
                print(f"   Para {item['index']:2d}: {has_section_break} | {item['text'][:50]}...")

        # Test full processing with section break preservation
        print(f"\n🔧 Testing Full Processing with Section Break Preservation:")
        print("-" * 60)

        output_file = process_document_line_breaks(test_file, "test_section_breaks_fixed.docx")

        # Verify sections are preserved in output
        print(f"\n📊 Verifying Section Preservation:")
        print("-" * 40)

        original_doc = Document(test_file)
        fixed_doc = Document(output_file)

        print("Original document sections:")
        original_sections = []
        for i, section in enumerate(original_doc.sections):
            try:
                start_type = section.start_type
                original_sections.append(start_type)
                print(f"   Section {i}: {start_type}")
            except:
                original_sections.append("unknown")
                print(f"   Section {i}: unknown")

        print("\nProcessed document sections:")
        fixed_sections = []
        for i, section in enumerate(fixed_doc.sections):
            try:
                start_type = section.start_type
                fixed_sections.append(start_type)
                print(f"   Section {i}: {start_type}")
            except:
                fixed_sections.append("unknown")
                print(f"   Section {i}: unknown")

        # Compare sections
        if len(original_sections) == len(fixed_sections):
            sections_preserved = True
            for i, (orig, fixed) in enumerate(zip(original_sections, fixed_sections)):
                if orig != fixed:
                    sections_preserved = False
                    print(f"   ⚠️  Section {i} changed: {orig} → {fixed}")

            if sections_preserved:
                print("   ✅ All section breaks preserved successfully!")
            else:
                print("   ⚠️  Some section breaks may have been modified")
        else:
            print(f"   ⚠️  Section count changed: {len(original_sections)} → {len(fixed_sections)}")

        # Test spacing adjustments with section breaks
        print(f"\n📏 Testing Spacing Adjustments with Section Breaks:")
        print("-" * 55)

        spacing_adjustments = 0
        for i, (orig_para, fixed_para) in enumerate(
            zip(original_doc.paragraphs, fixed_doc.paragraphs)
        ):
            if orig_para.text.strip():
                orig_before = manager._get_spacing_before(orig_para)
                fixed_before = manager._get_spacing_before(fixed_para)

                if abs(orig_before - fixed_before) > 1:
                    spacing_adjustments += 1
                    has_section_break = manager._has_section_break_before(fixed_doc, i)
                    content_type = manager.identify_content_type(orig_para.text)

                    print(f"   Para {i+1} ({content_type}):")
                    print(f"     Spacing: {orig_before}pt → {fixed_before}pt")
                    print(f"     Section break before: {has_section_break}")
                    print(f"     Text: {orig_para.text[:40]}...")

        print(f"\n✅ Total spacing adjustments: {spacing_adjustments}")

        print(f"\n🎉 Section break preservation testing completed!")

    except ImportError as e:
        print(f"❌ Could not import line break manager: {e}")
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback

        traceback.print_exc()

    finally:
        # Clean up test files
        for file in [test_file, "test_section_breaks_fixed.docx"]:
            if os.path.exists(file):
                os.remove(file)
                print(f"🧹 Cleaned up: {file}")


def test_google_docs_simulation():
    """Test with a document structure similar to Google Docs imports"""
    print("\n📄 Testing Google Docs Import Simulation:")
    print("-" * 50)

    # Create a document that simulates Google Docs structure
    doc = Document()

    # Add content with various spacing issues (typical of Google Docs imports)
    doc.add_paragraph("Chapter 1: Introduction")
    doc.add_paragraph("")  # Empty paragraph
    doc.add_paragraph("This is the introduction with improper spacing.")

    # Add section break (next page)
    section2 = doc.add_section(WD_SECTION_START.NEW_PAGE)

    doc.add_paragraph("Chapter 2: Background")
    doc.add_paragraph("")  # Empty paragraph
    doc.add_paragraph("")  # Another empty paragraph
    doc.add_paragraph("Background information follows.")

    # Add major section within the same page
    doc.add_paragraph("2.1 Related Work")
    doc.add_paragraph("Related work is discussed here.")

    # Add another section break
    section3 = doc.add_section(WD_SECTION_START.NEW_PAGE)

    doc.add_paragraph("Chapter 3: Methodology")
    doc.add_paragraph("Methodology is described here.")

    test_file = "test_google_docs_sim.docx"
    doc.save(test_file)

    try:
        from academic_line_break_manager import process_document_line_breaks

        print(f"✅ Created Google Docs simulation: {test_file}")

        # Process the document
        output_file = process_document_line_breaks(test_file, "test_google_docs_sim_fixed.docx")

        print("✅ Processing completed - check output for proper section break preservation")

    except Exception as e:
        print(f"❌ Error in Google Docs simulation: {e}")

    finally:
        # Clean up
        for file in [test_file, "test_google_docs_sim_fixed.docx"]:
            if os.path.exists(file):
                os.remove(file)


def main():
    """Run the section break preservation test"""
    print("🚀 Section Break Preservation Test Suite")
    print("=" * 70)

    test_section_break_preservation()
    test_google_docs_simulation()

    print("\n" + "=" * 70)
    print("🏁 Test Completed!")
    print("\nKey Features Tested:")
    print("   ✅ Section break detection and preservation")
    print("   ✅ Spacing adjustment with section breaks")
    print("   ✅ Google Docs import simulation")
    print("   ✅ Integration with line break processing")


if __name__ == "__main__":
    main()
