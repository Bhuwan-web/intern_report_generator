#!/usr/bin/env python3
"""
Test script for academic line break manager
Tests proper line break formatting according to academic standards
"""

import sys
import os
from docx import Document
from docx.shared import Pt

# Add src directory to path
sys.path.append("src")


def create_test_document_with_line_break_issues():
    """Create a test document with various line break issues"""
    doc = Document()

    # Add content with improper line breaks
    test_content = [
        # Chapter with no proper spacing
        ("Chapter 1: Introduction", "chapter_heading"),
        ("This is the introduction paragraph immediately after chapter.", "paragraph"),
        # Empty paragraphs (excessive)
        ("", "empty"),
        ("", "empty"),
        ("", "empty"),
        # Major section with improper spacing
        ("1. Literature Review", "major_section"),
        ("This section reviews existing literature without proper spacing.", "paragraph"),
        # Section heading with no spacing
        ("1.1 Background", "section_heading"),
        ("Background information follows immediately.", "paragraph"),
        # Subsection with improper spacing
        ("1.1.1 Problem Statement", "subsection_heading"),
        ("The problem statement is described here.", "paragraph"),
        # More excessive empty paragraphs
        ("", "empty"),
        ("", "empty"),
        # Another chapter
        ("Chapter 2: Methodology", "chapter_heading"),
        ("", "empty"),
        ("", "empty"),
        ("This chapter describes the methodology used.", "paragraph"),
        # Figure with no proper spacing
        ("Figure 2.1: System Architecture", "figure_table"),
        ("The figure shows the system architecture.", "paragraph"),
        # Table with improper spacing
        ("Table 2.1: Comparison Results", "figure_table"),
        ("The table presents comparison results.", "paragraph"),
        # Major section
        ("2. Implementation", "major_section"),
        ("Implementation details are provided here.", "paragraph"),
        # Nested sections
        ("2.1 System Design", "section_heading"),
        ("System design considerations.", "paragraph"),
        ("2.1.1 Database Design", "subsection_heading"),
        ("Database design details.", "paragraph"),
        ("2.1.2 User Interface Design", "subsection_heading"),
        ("User interface design details.", "paragraph"),
        # Block quote (indented)
        (
            '    "This is a block quote that should have proper spacing before and after."',
            "quote_block",
        ),
        ("Regular paragraph after quote.", "paragraph"),
        # List items
        ("• First bullet point", "list_item"),
        ("• Second bullet point", "list_item"),
        ("• Third bullet point", "list_item"),
        # Final paragraph
        ("This is the final paragraph of the document.", "paragraph"),
    ]

    for text, content_type in test_content:
        para = doc.add_paragraph(text)

        # Add some random spacing to simulate real document issues
        if content_type == "chapter_heading":
            para.paragraph_format.space_before = Pt(0)  # Should be more
            para.paragraph_format.space_after = Pt(0)  # Should be more
        elif content_type == "major_section":
            para.paragraph_format.space_before = Pt(6)  # Should be more
            para.paragraph_format.space_after = Pt(0)  # Should be more
        elif content_type == "section_heading":
            para.paragraph_format.space_before = Pt(0)  # Should be some
            para.paragraph_format.space_after = Pt(12)  # Should be less

    test_file = "test_line_breaks.docx"
    doc.save(test_file)
    return test_file


def test_line_break_manager():
    """Test the academic line break manager"""
    print("📏 Testing Academic Line Break Manager")
    print("=" * 60)

    # Create test document
    test_file = create_test_document_with_line_break_issues()
    print(f"✅ Created test document: {test_file}")

    try:
        from academic_line_break_manager import (
            AcademicLineBreakManager,
            process_document_line_breaks,
        )

        # Test content type identification
        print("\n🔍 Testing Content Type Identification:")
        print("-" * 50)

        manager = AcademicLineBreakManager()

        test_texts = [
            "Chapter 1: Introduction",
            "CHAPTER 2: METHODOLOGY",
            "1. Literature Review",
            "2. Implementation",
            "1.1 Background",
            "2.1 System Design",
            "1.1.1 Problem Statement",
            "2.1.2 Database Design",
            "Figure 2.1: System Architecture",
            "Table 3.1: Results",
            "This is a regular paragraph.",
            "• First bullet point",
            "1. This could be a list item",
            '    "This is a block quote."',
        ]

        for text in test_texts:
            content_type = manager.identify_content_type(text)
            print(f"   '{text}' → {content_type}")

        # Test document structure analysis
        print(f"\n📊 Testing Document Structure Analysis:")
        print("-" * 50)

        doc = Document(test_file)
        structure = manager.analyze_document_structure(doc)

        print("Document structure identified:")
        for item in structure[:10]:  # Show first 10 items
            if item["text"]:
                print(
                    f"   {item['index']:2d}. {item['content_type']:15s} | {item['text'][:50]}..."
                )

        if len(structure) > 10:
            print(f"   ... and {len(structure) - 10} more items")

        # Test full line break processing
        print(f"\n🔧 Testing Full Line Break Processing:")
        print("-" * 50)

        output_file = process_document_line_breaks(test_file, "test_line_breaks_fixed.docx")

        # Compare before and after
        print(f"\n📊 Before/After Analysis:")
        print("-" * 30)

        original_doc = Document(test_file)
        fixed_doc = Document(output_file)

        # Analyze spacing changes
        spacing_changes = 0
        for i, (orig_para, fixed_para) in enumerate(
            zip(original_doc.paragraphs, fixed_doc.paragraphs)
        ):
            if orig_para.text.strip():
                orig_before = manager._get_spacing_before(orig_para)
                orig_after = manager._get_spacing_after(orig_para)
                fixed_before = manager._get_spacing_before(fixed_para)
                fixed_after = manager._get_spacing_after(fixed_para)

                if abs(orig_before - fixed_before) > 1 or abs(orig_after - fixed_after) > 1:
                    spacing_changes += 1
                    content_type = manager.identify_content_type(orig_para.text)
                    print(f"   Para {i+1} ({content_type}):")
                    print(f"     Before: {orig_before}pt → {fixed_before}pt")
                    print(f"     After:  {orig_after}pt → {fixed_after}pt")
                    print(f"     Text: {orig_para.text[:40]}...")

        print(f"\n✅ Total spacing adjustments made: {spacing_changes}")

        # Test line break standards
        print(f"\n📋 Testing Line Break Standards:")
        print("-" * 40)

        print("Applied standards:")
        for content_type, standards in manager.line_break_standards.items():
            print(f"   {content_type}:")
            print(f"     Before: {standards['before']} line breaks")
            print(f"     After:  {standards['after']} line breaks")
            print(f"     Description: {standards['description']}")

        print(f"\n🎉 Line break manager testing completed!")

    except ImportError as e:
        print(f"❌ Could not import line break manager: {e}")
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback

        traceback.print_exc()

    finally:
        # Clean up test files
        for file in [test_file, "test_line_breaks_fixed.docx"]:
            if os.path.exists(file):
                os.remove(file)
                print(f"🧹 Cleaned up: {file}")


def test_integration_with_main():
    """Test integration with main processing pipeline"""
    print("\n🔗 Testing Integration with Main Pipeline:")
    print("-" * 50)

    # Create a simple test document
    doc = Document()
    doc.add_paragraph("Chapter 1: Introduction")
    doc.add_paragraph("This is the introduction.")
    doc.add_paragraph("")  # Empty paragraph
    doc.add_paragraph("")  # Another empty paragraph
    doc.add_paragraph("1. Literature Review")
    doc.add_paragraph("Literature review content.")

    test_file = "test_integration.docx"
    doc.save(test_file)

    try:
        # Test using main.py approach
        import subprocess
        import sys

        # Test the linebreaks action
        result = subprocess.run(
            [sys.executable, "src/main.py", test_file, "linebreaks"],
            capture_output=True,
            text=True,
            cwd=".",
        )

        if result.returncode == 0:
            print("✅ Integration with main.py successful")
            print("Output preview:")
            print(result.stdout[-300:])  # Show last 300 characters
        else:
            print("❌ Integration failed")
            print("Error:", result.stderr)

    except Exception as e:
        print(f"❌ Integration test failed: {e}")

    finally:
        # Clean up
        for file in [test_file, f"line_break_fixed_{test_file.replace('.docx', '')}.docx"]:
            if os.path.exists(file):
                os.remove(file)


def main():
    """Run the line break manager test"""
    print("🚀 Academic Line Break Manager Test Suite")
    print("=" * 70)

    test_line_break_manager()
    test_integration_with_main()

    print("\n" + "=" * 70)
    print("🏁 Test Completed!")
    print("\nTo use line break manager:")
    print("   python src/academic_line_break_manager.py your_document.docx")
    print("   python src/main.py your_document.docx linebreaks")
    print("   python src/main.py your_document.docx complete")


if __name__ == "__main__":
    main()
